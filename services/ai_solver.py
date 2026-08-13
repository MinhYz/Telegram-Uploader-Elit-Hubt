import os
from pathlib import Path
from typing import List, Tuple
from google import genai
from pypdf import PdfReader
from config.settings import GEMINI_API_KEY, DOWNLOAD_DIR
from utils.logger import logger

class AISolverService:
    def __init__(self):
        self.api_key = GEMINI_API_KEY
        self.client = genai.Client(api_key=self.api_key) if self.api_key else None

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        try:
            reader = PdfReader(pdf_path)
            extracted = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted.append(text)
            return "\n".join(extracted)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {e}")
            return ""

    async def solve_assignment_file(self, file_paths: List[str], assignment_id: str, student_info: dict) -> Tuple[bool, Path, str]:
        """Process assignment file via OCR/Gemini API and output solved Excel/Word documents."""
        if not self.client:
            return False, Path(""), "⚠️ CHƯA CẤU HÌNH GEMINI_API_KEY trong file .env!"

        combined_prompt_text = []
        for fp in file_paths:
            path = Path(fp)
            if path.suffix.lower() == ".pdf":
                pdf_text = self.extract_text_from_pdf(path)
                combined_prompt_text.append(f"--- Nội dung PDF {path.name} ---\n{pdf_text}")

        prompt = (
            f"Bạn là chuyên gia giải bài tập CNTT/Kinh tế tại HUBT.\n"
            f"Thông tin sinh viên: {student_info.get('name', 'Trần Tuấn Minh')} - MSV: {student_info.get('id', '16T-Tin3')}\n\n"
            f"Nội dung đề bài:\n" + "\n".join(combined_prompt_text) + "\n\n"
            f"Hãy đưa ra lời giải chi tiết, chuẩn xác, kèm công thức tính toán theo chuẩn HUBT."
        )

        try:
            response = self.client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt,
            )
            solution_text = response.text or "Không có phản hồi từ Gemini."

            # Save detailed docx solution
            doc_path = DOWNLOAD_DIR / f"Loi_Giai_Assignment_{assignment_id}.docx"
            from docx import Document
            doc = Document()
            doc.add_heading(f"Lời Giải Bài Tập #{assignment_id}", 0)
            doc.add_paragraph(solution_text)
            doc.save(doc_path)

            caption = f"✨ **ĐÃ GIẢI XONG BÀI TẬP #{assignment_id} BẰNG GEMINI AI**\n📎 File lời giải chi tiết: `{doc_path.name}`"
            return True, doc_path, caption
        except Exception as e:
            logger.error(f"Gemini AI solve error: {e}")
            return False, Path(""), f"❌ Lỗi giải bài AI: {str(e)}"

    async def paraphrase_anti_plagiarism(self, text: str) -> str:
        """Paraphrase output to ensure unique solutions across multiple students."""
        if not self.client:
            return text
        try:
            response = self.client.models.generate_content(
                model="gemini-2.5-flash",
                contents=f"Viết lại đoạn văn sau theo cách diễn đạt khác nhưng giữ nguyên ý nghĩa và độ chính xác để chống đạo văn:\n\n{text}",
            )
            return response.text or text
        except Exception as e:
            logger.error(f"Paraphrase error: {e}")
            return text

ai_solver = AISolverService()
