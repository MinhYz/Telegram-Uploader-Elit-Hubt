from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config.settings import DOWNLOAD_DIR
from utils.logger import logger

class HUBTCoverPageGenerator:
    """Generate standardized HUBT Assignment Cover Pages (Times New Roman 13, 2cm-2.5cm margins)."""

    @staticmethod
    def generate_cover_docx(
        course_name: str,
        assignment_title: str,
        student_name: str = "Trần Tuấn Minh",
        student_id: str = "16T-Tin3",
        lecturer_name: str = "Giảng viên HUBT",
        output_path: Path = None,
    ) -> Path:
        if not output_path:
            output_path = DOWNLOAD_DIR / f"Bia_Bai_Tap_{student_id}.docx"

        doc = Document()

        # Set Margins (Top: 2cm, Bottom: 2cm, Left: 2.5cm, Right: 2cm)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.78)    # 2.0 cm
            section.bottom_margin = Inches(0.78) # 2.0 cm
            section.left_margin = Inches(0.98)   # 2.5 cm
            section.right_margin = Inches(0.78)  # 2.0 cm

        # Header Title
        p_top = doc.add_paragraph()
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_top = p_top.add_run("TRƯỜNG ĐẠI HỌC KINH DOANH VÀ CÔNG NGHỆ HÀ NỘI\nKHOA CÔNG NGHỆ THÔNG TIN\n------------------***------------------")
        run_top.font.name = "Times New Roman"
        run_top.font.size = Pt(13)
        run_top.font.bold = True

        doc.add_paragraph("\n\n")

        # Main Title
        p_main = doc.add_paragraph()
        p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_main = p_main.add_run(f"BÁO CÁO BÀI TẬP LỚP / TIỂU LUẬN\n\n{assignment_title.upper()}")
        run_main.font.name = "Times New Roman"
        run_main.font.size = Pt(18)
        run_main.font.bold = True
        run_main.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph("\n\n\n")

        # Info Section
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        info_text = (
            f"   Môn học: {course_name}\n"
            f"   Giảng viên hướng dẫn: {lecturer_name}\n"
            f"   Sinh viên thực hiện: {student_name}\n"
            f"   Mã sinh viên / Lớp: {student_id}\n"
        )
        run_info = p_info.add_run(info_text)
        run_info.font.name = "Times New Roman"
        run_info.font.size = Pt(13)

        doc.add_paragraph("\n\n\n")

        # Footer Date
        p_foot = doc.add_paragraph()
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_foot = p_foot.add_run("HÀ NỘI - 2026")
        run_foot.font.name = "Times New Roman"
        run_foot.font.size = Pt(13)
        run_foot.font.italic = True

        doc.save(output_path)
        logger.info(f"Generated HUBT cover page docx at {output_path}")
        return output_path

pdf_cover_generator = HUBTCoverPageGenerator()
